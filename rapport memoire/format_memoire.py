# -*- coding: utf-8 -*-
"""
Mise en forme automatique du mémoire Master BlockTask.

Passes :
  1. Styles de titres (Titre1=chapitres, Titre2=sections X.Y, Titre3=sous-sections X.Y.Z)
  2. Sauts de page avant chaque chapitre
  3. Listes automatiques des figures / tableaux (champs TOC)
  4. Légendes numérotées par chapitre (champs SEQ) + paragraphe explicatif sous chaque tableau
  5. Suppression des blocs de paragraphes vides (pages blanches)
  6. Justification du texte du corps
  7. Pagination (Page X / Y)
"""
import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

SRC = Path(r"C:\Users\PC\Documents\blocktask_myself\rapport memoire\memoire de fin detude Master.docx")
DST = SRC.with_name("memoire de fin detude Master - FORMATE v6.docx")

# Termes du glossaire -> numero (ordre du glossaire) + regex de recherche
GLOSSARY_TERMS = [
    (1, "BlockTask", r"BlockTask"),
    (2, "Escrow", r"escrow"),
    (3, "FCFA / XOF", r"FCFA"),
    (4, "Gig economy", r"gig\s+economy"),
    (5, "Hybrid architecture", r"architecture\s+hybride"),
    (6, "KYC", r"KYC"),
    (7, "Mobile Money", r"Mobile\s+Money"),
    (8, "NINA", r"NINA"),
    (9, "Oracle", r"oracles?"),
    (10, "Smart contract", r"smart\s+contracts?"),
    (11, "UEMOA", r"UEMOA"),
]

# Identifiants de style RÉELS du document (vérifiés dans styles.xml)
HEADING1 = "Titre1"
HEADING2 = "Titre2"
HEADING3 = "Titre3"
CAPTION_STYLE = "Lgende"
NORMAL = "Normal"

TABLE_TITLES = {
    3: "Modèle hybride BlockTask — Couches et technologies",
    4: "Objectifs spécifiques et résultats attendus",
    5: "Outils et technologies de développement",
    6: "Organisation du mémoire — Chapitres et contenu",
    7: "Comparaison des plateformes centralisées (Upwork, Fiverr, TaskRabbit)",
    8: "Vulnérabilités des systèmes de réputation numériques",
    9: "Apports de la blockchain pour la délégation de tâches",
    10: "Limites des systèmes blockchain",
    11: "Comparaison synthétique — Solutions existantes vs BlockTask",
    12: "Modèle hybride BlockTask — Couches implémentées",
    13: "Périmètre fonctionnel BlockTask",
    14: "Acteurs du système",
    15: "États du cycle de vie d'une mission",
    16: "Règles de gestion (extrait)",
    17: "Exigences fonctionnelles principales",
    18: "Exigences non fonctionnelles",
    19: "Synthèse des cas d'utilisation",
    20: "Scénario 1 — Inscription, KYC et première connexion",
    21: "Scénario 2 — Création de mission et paiement escrow",
    22: "Scénario 3 — Candidature, acceptation et caution",
    23: "Scénario 4 — Exécution, preuves et suivi GPS",
    24: "Scénario 5 — Validation client et libération des fonds",
    25: "Scénario 6 — Litige et arbitrage administrateur",
    26: "Scénario 7 — Gestion entreprise (B2B)",
    27: "Matrice de traçabilité conception → implémentation",
    28: "Composants matériels et logiciels — Versions minimales",
    29: "Outils utilisés par couche",
    30: "Architecture 3-tiers — Couches et responsabilités",
    31: "Modèle hybride Mali — Flux et implémentation",
    32: "Langages de programmation utilisés",
    33: "Bibliothèques backend (extrait requirements.txt)",
    34: "Bibliothèques frontend (extrait package.json)",
    35: "Bibliothèques smart contracts",
    36: "Modules backend Django",
    37: "Actions API du workflow mission",
    38: "Fonctions EscrowContract Solidity",
    39: "API d'enregistrement blockchain",
    40: "Pages administrateur Angular",
    41: "Mesures de sécurité — Authentification",
    42: "Protection des communications et données",
    43: "Mécanismes de sécurité des smart contracts",
    44: "Services Docker Compose",
    45: "Configuration blockchain par environnement",
    46: "Difficultés d'implémentation et solutions adoptées",
    47: "Stratégie de test multi-niveaux",
    48: "Environnement de test",
    49: "Couverture des exigences fonctionnelles",
    50: "Couverture des exigences non fonctionnelles",
    51: "Fixtures communes (conftest.py)",
    52: "Tests du flux MVP (pytest)",
    53: "Tests KYC et accès plateforme",
    54: "Tests rôles doubles",
    55: "Synthèse tests backend",
    56: "Tests EscrowContract (Hardhat)",
    57: "Tests ReputationContract (Hardhat)",
    58: "Tests LitigationContract (Hardhat)",
    59: "Scénario manuel SC1 — Inscription, KYC et connexion",
    60: "Scénario manuel SC2 — Création mission et paiement escrow",
    61: "Scénario manuel SC3 — Candidature, acceptation et caution",
    62: "Scénario manuel SC4 — Exécution, preuves et GPS",
    63: "Scénario manuel SC5 — Validation et paiement prestataire",
    64: "Scénario manuel SC6 — Litige et arbitrage",
    65: "Scénario manuel SC7 — Administration plateforme",
    66: "Checklist sécurité",
    67: "Synthèse quantitative des tests",
    68: "Couverture par objectif spécifique (OS1–OS10)",
    69: "Lacunes identifiées dans la couverture de tests",
    70: "Modèle hybride — Couches et résultats observés",
    71: "Bilan d'atteinte des objectifs spécifiques",
    72: "Validation des hypothèses de travail",
    73: "BlockTask vs alternatives — Synthèse post-implémentation",
    74: "Synthèse des réalisations du projet",
    75: "Difficultés rencontrées et solutions apportées",
    76: "Glossaire — Définitions des termes principaux",
}

FIGURE_TITLES = [
    "Capture d'écran — Plateforme Upwork",
    "Capture d'écran — Interface Upwork (détail)",
    "Capture d'écran — Plateforme Fiverr",
    "Capture d'écran — Plateforme TaskRabbit",
    "Diagramme d'états — Cycle de vie d'une mission",
    "Diagramme de contexte C4 — BlockTask",
    "Diagramme de conteneurs C4 — Architecture BlockTask",
    "Diagramme de composants — Backend Django",
    "Diagramme de cas d'utilisation — BlockTask",
    "Diagramme de classes — Cœur métier",
    "Diagramme de séquence — Création mission et paiement FCFA",
    "Diagramme de séquence — Caution, exécution et validation",
    "Diagramme d'activité — Processus global mission",
    "Diagramme de déploiement — Environnement prototype",
    "Architecture globale du projet BlockTask",
    "Structure du frontend Angular — Organisation par features",
    "Arborescence des smart contracts Hardhat",
    "Difficultés d'implémentation — Schéma récapitulatif",
    "Capture — Résultats des tests automatisés",
]

# --- Captures d'écran de l'application (insérées en fin de Chapitre 4) ---
ASSET_DIR = Path(
    r"C:\Users\PC\.cursor\projects\c-Users-PC-Documents-blocktask-myself\assets"
)
SS_PREFIX = (
    "c__Users_PC_AppData_Roaming_Cursor_User_workspaceStorage_"
    "8c3f9271a254aa41fdfd7517f4e00799_images_Screenshot_2026-06-28_"
)

# Ordre logique : visiteur -> client -> prestataire -> administrateur
# ("group", libellé) ou ("fig", suffixe-fichier, légende)
SCREENSHOTS = [
    ("group", "Espace visiteur"),
    ("fig", "064744-3d56f651-1068-477f-bb0c-cb8457808bf5.png",
     "Page d'accueil publique de BlockTask et recherche de services"),
    ("fig", "064801-605a9077-e11d-46a5-8fc3-14e05454a254.png",
     "Page « Comment ça marche » : le parcours en quatre étapes côté client"),
    ("fig", "064710-a15e0d2a-f06b-4208-9072-928780a07289.png",
     "Inscription en trois étapes avec choix du profil (client, prestataire, entreprise)"),
    ("fig", "064732-e6b0a6b3-505b-44a1-afcf-a8313c88fd5b.png",
     "Page de connexion à l'espace BlockTask"),

    ("group", "Espace client"),
    ("fig", "064829-d0093652-6d59-43c1-9673-573080f7e424.png",
     "Tableau de bord du client et menu de changement d'espace"),
    ("fig", "064844-1e1c8242-1240-4f00-8fa1-8a2df5b8f1f6.png",
     "Création d'une mission via l'assistant en quatre étapes"),
    ("fig", "064900-964ffbc2-af46-4e8d-8190-337dd10b2de2.png",
     "Attribution directe d'une mission à un prestataire"),
    ("fig", "065047-14b89cb6-9e50-463a-afb5-09dc8e23196b.png",
     "Centre de notifications de l'utilisateur"),
    ("fig", "064934-43b5ee6a-0f68-4eb2-8b23-7f96b138c55f.png",
     "Profil utilisateur et vérification d'identité (NINA et téléphone)"),

    ("group", "Espace prestataire"),
    ("fig", "064952-0a3c2f07-0af2-4034-ad0c-1fcb7a6308fb.png",
     "Tableau de bord prestataire — niveau de réputation Bronze"),
    ("fig", "065104-cd6b35ee-52b0-4127-bf73-9bf2d14ff76e.png",
     "Tableau de bord prestataire — revenus et niveau Silver"),
    ("fig", "065121-1f3da4da-fb35-4b9c-9bf9-02d0f81aae25.png",
     "Liste des missions assignées au prestataire"),
    ("fig", "065157-498372a2-b384-4313-9f17-16a745b48434.png",
     "Détail d'une mission assignée et suivi des étapes"),
    ("fig", "065211-8b21a1dc-3287-4604-ab87-acfd3fdf35e3.png",
     "Itinéraire, sécurisation escrow et soumission des preuves d'exécution"),

    ("group", "Espace administrateur"),
    ("fig", "065244-182ad4ae-9163-43f4-9c1a-6b2e9b57ed18.png",
     "Tableau de bord de l'administrateur"),
    ("fig", "065257-c772ef27-e68f-4828-a236-483805ee3e06.png",
     "Administration — gestion des utilisateurs"),
    ("fig", "065311-e56b501f-116b-46a0-8139-bc7cc2fe398c.png",
     "Administration — supervision des missions"),
    ("fig", "065335-3f4a6f4e-2ad0-40a8-a1c6-f23a1c61e1f2.png",
     "Administration — validation des dossiers KYC"),
    ("fig", "065409-4ccabdc8-83ef-4fb1-a3e0-6304618d5f88.png",
     "Administration — paramètres de la plateforme"),
    ("fig", "065354-b5d63d9c-2925-4930-93ba-e126cdaa16bb.png",
     "Administration — transactions blockchain et escrow"),
]


RE_OLD_CAPTION = re.compile(
    r"^(Figure|Tableau)\s+\d+([\.\-,]\d+)?(\s*—\s*.+)?$", re.IGNORECASE
)
RE_H2 = re.compile(r"^(\d+\.\d+)\s+\S")
RE_H3 = re.compile(r"^(\d+\.\d+\.\d+)\s+\S")
RE_MANUAL_TABLE_CAP = re.compile(r"^Tableau\s+\d+\.\d+\s*—", re.IGNORECASE)
RE_MANUAL_FIG_CAP = re.compile(r"^Figure\s+\d+\.\d+\s*—", re.IGNORECASE)


def w_tag(tag):
    return qn(f"w:{tag}")


def make_run(text="", bold=False, italic=False):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    if len(rPr):
        r.append(rPr)
    if text:
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        r.append(t)
    return r


def make_field(instr, placeholder="1"):
    runs = []
    r = OxmlElement("w:r")
    fc = OxmlElement("w:fldChar")
    fc.set(w_tag("fldCharType"), "begin")
    r.append(fc)
    runs.append(r)

    r_instr = OxmlElement("w:r")
    instr_el = OxmlElement("w:instrText")
    instr_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_el.text = f" {instr.strip()} "
    r_instr.append(instr_el)
    runs.append(r_instr)

    r_sep = OxmlElement("w:r")
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(w_tag("fldCharType"), "separate")
    r_sep.append(fc_sep)
    runs.append(r_sep)

    runs.append(make_run(placeholder))

    r_end = OxmlElement("w:r")
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(w_tag("fldCharType"), "end")
    r_end.append(fc_end)
    runs.append(r_end)
    return runs


def make_superscript_run(num):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    va = OxmlElement("w:vertAlign")
    va.set(w_tag("val"), "superscript")
    rPr.append(va)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = str(num)
    r.append(t)
    return r


def make_caption_paragraph(label, seq_name, title, chapter, placeholder):
    """Légende numérotée par chapitre : « Label C.N — titre »
    chapter : préfixe littéral (ex. '3', 'G')
    N : champ SEQ qui se réinitialise à chaque Titre1 (\\s 1)
    """
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(w_tag("val"), CAPTION_STYLE)
    pPr.append(pStyle)
    jc = OxmlElement("w:jc")
    jc.set(w_tag("val"), "center")
    pPr.append(jc)
    p.append(pPr)

    p.append(make_run(f"{label} ", bold=True))
    p.append(make_run(f"{chapter}.", bold=True))
    for r in make_field(f"SEQ {seq_name} \\* ARABIC \\s 1", str(placeholder)):
        p.append(r)
    if title:
        p.append(make_run(f" — {title}", italic=True))
    return p


SKIP_HEADERS = {
    "n°", "no", "n", "id", "#", "uc", "code", "réf", "ref", "n0", "nº",
}


def _cell_text(tc):
    parts = []
    for p in tc.findall(".//" + w_tag("p")):
        t = "".join(x.text or "" for x in p.findall(".//" + w_tag("t")))
        t = t.strip()
        if t:
            parts.append(t)
    return " ".join(parts)


def _is_codey(vals):
    """Colonne faite de numéros/codes courts (N°, 1, 2, A1...) -> peu parlante."""
    if not vals:
        return True
    c = sum(1 for v in vals if re.fullmatch(r"[\d.\-–—/]+", v) or len(v) <= 3)
    return c / len(vals) > 0.6


def _table_rows(tbl):
    trs = tbl.findall(w_tag("tr"))
    if len(trs) == 1:
        tcs = trs[0].findall(w_tag("tc"))
        if len(tcs) == 1:
            inner = tcs[0].find(w_tag("tbl"))
            if inner is not None:
                return _table_rows(inner)
    rows = []
    for tr in trs:
        rows.append([_cell_text(tc) for tc in tr.findall(w_tag("tc"))])
    return rows


def _french_join(items):
    items = [i for i in items if i]
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    return ", ".join(items[:-1]) + " et " + items[-1]


def _french_enum(items, limit=10):
    if len(items) <= limit:
        return _french_join(items)
    return ", ".join(items[:limit]) + ", etc."


def _pluralize(word):
    w = word.strip()
    if not w or w.endswith(("s", "x", "z")):
        return w
    return w + "s"


def _lower_first(s):
    s = s.strip()
    if not s:
        return s
    if s[:2].isupper():
        return s
    return s[0].lower() + s[1:]


def _dedupe(seq):
    seen = set()
    out = []
    for x in seq:
        if x and x not in seen:
            seen.add(x)
            out.append(x)
    return out


def make_table_explanation(tbl_elem, title):
    """Paragraphe explicatif de fond : énonce le contenu réel du tableau
    (éléments couverts) pour le comprendre sans avoir à le lire."""
    rows = _table_rows(tbl_elem)
    headers = rows[0] if rows else []
    data = rows[1:] if len(rows) > 1 else []
    ncols = len(headers)

    # Indices candidats = en-têtes non triviales (pas N°/ID...)
    candidates = [
        i for i, h in enumerate(headers) if h.strip().lower() not in SKIP_HEADERS
    ] or list(range(ncols))

    # Choisir une colonne « libellé » porteuse de sens (pas une colonne de codes)
    label_idx = candidates[0]
    for i in candidates:
        vals = [r[i].strip() for r in data if len(r) > i and r[i].strip()]
        if not _is_codey(vals):
            label_idx = i
            break

    items = _dedupe(
        r[label_idx].strip()
        for r in data
        if len(r) > label_idx and r[label_idx].strip()
        and not re.fullmatch(r"[\d.\-–—/\s]+", r[label_idx].strip())
    )
    subj = (headers[label_idx].strip().lower() if headers else "élément") or "élément"
    others = [
        headers[i].strip().lower()
        for i in range(ncols)
        if i != label_idx
        and headers[i].strip()
        and headers[i].strip().lower() not in SKIP_HEADERS
    ]
    n = len(items)
    avglen = (sum(len(x) for x in items) / n) if n else 0

    phrase = f"Le tableau « {title} » "
    if n >= 1 and avglen <= 34:
        listing = _french_enum(items, limit=10)
        if others:
            phrase += (
                f"passe en revue {n} {_pluralize(subj)} — {listing} — "
                f"en précisant à chaque fois {_french_join(others)}."
            )
        else:
            phrase += f"énumère {n} {_pluralize(subj)} : {listing}."
    else:
        if others:
            phrase += (
                f"détaille, pour chaque {subj}, {_french_join(others)}, "
                "afin d'en restituer le contenu sans avoir à le parcourir ligne à ligne."
            )
        else:
            phrase += "récapitule les éléments essentiels abordés dans cette section."

    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(w_tag("val"), NORMAL)
    pPr.append(pStyle)
    jc = OxmlElement("w:jc")
    jc.set(w_tag("val"), "both")
    pPr.append(jc)
    p.append(pPr)
    p.append(make_run(phrase))
    return p


def make_toc_paragraph(label):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(w_tag("val"), NORMAL)
    pPr.append(pStyle)
    p.append(pPr)
    for r in make_field(
        f'TOC \\h \\z \\c "{label}"',
        f"[Liste des {label.lower()}s — selectionner puis F9 pour actualiser]",
    ):
        p.append(r)
    return p


def paragraph_text(p_elem):
    return "".join(t.text or "" for t in p_elem.findall(".//" + w_tag("t"))).strip()


def has_page_break(p_elem):
    return any(
        br.get(w_tag("type")) == "page" for br in p_elem.findall(".//" + w_tag("br"))
    )


def has_image(p_elem):
    return bool(
        p_elem.findall(".//" + w_tag("drawing")) or p_elem.findall(".//" + w_tag("pict"))
    )


def get_style(p_elem):
    pPr = p_elem.find(w_tag("pPr"))
    if pPr is None:
        return NORMAL
    pStyle = pPr.find(w_tag("pStyle"))
    return pStyle.get(w_tag("val")) if pStyle is not None else NORMAL


def set_style(p_elem, style_id):
    pPr = p_elem.find(w_tag("pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    pStyle = pPr.find(w_tag("pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(w_tag("val"), style_id)


def set_jc(p_elem, val):
    pPr = p_elem.find(w_tag("pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    jc = pPr.find(w_tag("jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(w_tag("val"), val)


def strip_bold(p_elem):
    for r in p_elem.findall(w_tag("r")):
        rPr = r.find(w_tag("rPr"))
        if rPr is None:
            continue
        for tag in ("b", "bCs"):
            el = rPr.find(w_tag(tag))
            if el is not None:
                rPr.remove(el)


def insert_page_break_before(p_elem):
    if has_page_break(p_elem):
        return
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(w_tag("type"), "page")
    r.append(br)
    first_r = p_elem.find(w_tag("r"))
    if first_r is not None:
        first_r.addprevious(r)
    else:
        p_elem.append(r)


def remove_paragraph(p_elem):
    parent = p_elem.getparent()
    if parent is not None:
        parent.remove(p_elem)


def is_real_chapter_heading(txt):
    return bool(re.match(r"^CHAPITRE\s+\d+\s*:", txt, re.IGNORECASE))


# ---------------------------------------------------------------------------
# PASSES
# ---------------------------------------------------------------------------
def fix_heading_styles(body):
    stats = {"h1": 0, "h2": 0, "h3": 0, "pb": 0}
    for child in list(body.iterchildren()):
        if child.tag.split("}")[-1] != "p":
            continue
        txt = paragraph_text(child)
        if not txt:
            continue

        if is_real_chapter_heading(txt) or txt.upper().replace("\xa0", " ") in (
            "BIBLIOGRAPHIE",
            "GLOSSAIRE ET MOTS-CLÉS",
        ):
            set_style(child, HEADING1)
            strip_bold(child)
            stats["h1"] += 1
            m = re.match(r"^CHAPITRE\s+(\d+)\s*:", txt, re.IGNORECASE)
            if (m and int(m.group(1)) >= 2) or txt.upper().startswith("BIBLIOGRAPHIE"):
                insert_page_break_before(child)
                stats["pb"] += 1
            continue

        # Paragraphes descriptifs « Chapitre N — ... » du sommaire : ignorer
        if txt.startswith("Chapitre ") and " — " in txt:
            continue

        if RE_H3.match(txt):
            set_style(child, HEADING3)
            strip_bold(child)
            stats["h3"] += 1
        elif RE_H2.match(txt):
            set_style(child, HEADING2)
            strip_bold(child)
            stats["h2"] += 1
    return stats


def replace_manual_lists(body):
    replaced = 0
    children = list(body.iterchildren())
    for i, child in enumerate(children):
        if child.tag.split("}")[-1] != "p":
            continue
        txt = paragraph_text(child).upper().replace("\xa0", " ")
        label = None
        if "LISTE DES FIGURES" in txt:
            label = "Figure"
        elif "LISTE DES TABLEAUX" in txt:
            label = "Tableau"
        if not label:
            continue
        for j in range(i + 1, len(children)):
            if children[j].tag.split("}")[-1] == "tbl":
                remove_paragraph(children[j])
                replaced += 1
                break
        insert_after = make_toc_paragraph(label)
        child.addnext(insert_after)
    return replaced


def process_captions(body):
    figure_idx = 0
    table_idx = 0
    to_remove = []
    pending_fig_title = None
    pending_tbl_title = None
    current_chapter = "1"
    fig_counter = {}
    tbl_counter = {}

    for child in list(body.iterchildren()):
        tag = child.tag.split("}")[-1]

        if tag == "p":
            txt = paragraph_text(child)

            # Suivi du chapitre courant via les Titre1
            if get_style(child) == HEADING1:
                up = txt.upper().replace("\xa0", " ")
                m = re.match(r"^CHAPITRE\s+(\d+)", up)
                if m:
                    current_chapter = m.group(1)
                elif up.startswith("GLOSSAIRE"):
                    current_chapter = "G"
                elif up.startswith("BIBLIOGRAPHIE"):
                    current_chapter = "B"

            if (
                RE_OLD_CAPTION.match(txt)
                or txt in ("Tableau 1", "Tableau récapitulatif")
                or re.match(r"^Figure\s+\d+[,.]", txt, re.I)
            ):
                to_remove.append(child)
                continue
            if RE_MANUAL_FIG_CAP.match(txt):
                pending_fig_title = re.sub(r"^Figure\s+\d+\.\d+\s*—\s*", "", txt).strip()
                to_remove.append(child)
                continue
            if RE_MANUAL_TABLE_CAP.match(txt):
                pending_tbl_title = re.sub(r"^Tableau\s+\d+\.\d+\s*—\s*", "", txt).strip()
                to_remove.append(child)
                continue
            if "Figure" in txt and not has_image(child) and "—" in txt:
                m = re.search(r"—\s*Figure\s+\d+\.\d+\s*$", txt)
                if m:
                    cand = txt[: m.start()].strip()
                    if len(cand) < 80:
                        pending_fig_title = cand
                        to_remove.append(child)

            if has_image(child):
                title = pending_fig_title or (
                    FIGURE_TITLES[figure_idx]
                    if figure_idx < len(FIGURE_TITLES)
                    else f"Illustration {figure_idx + 1}"
                )
                pending_fig_title = None
                fig_counter[current_chapter] = fig_counter.get(current_chapter, 0) + 1
                ph = str(fig_counter[current_chapter])
                child.addnext(
                    make_caption_paragraph("Figure", "Figure", title, current_chapter, ph)
                )
                figure_idx += 1

        elif tag == "tbl":
            # 0 = abréviations, 1/2 = anciennes listes -> pas de légende
            if table_idx >= 3:
                title = pending_tbl_title or TABLE_TITLES.get(table_idx, "Tableau de données")
                pending_tbl_title = None
                tbl_counter[current_chapter] = tbl_counter.get(current_chapter, 0) + 1
                ph = str(tbl_counter[current_chapter])
                child.addprevious(
                    make_caption_paragraph("Tableau", "Tableau", title, current_chapter, ph)
                )
                child.addnext(make_table_explanation(child, title))
            table_idx += 1

    for p in to_remove:
        remove_paragraph(p)

    return {"figures": figure_idx, "tables": max(0, table_idx - 3)}


def collapse_blanks(body, keep=1):
    """Réduit toute série de paragraphes vides consécutifs à `keep` maximum."""
    removed = 0
    streak = []

    def flush():
        nonlocal removed
        for extra in streak[keep:]:
            remove_paragraph(extra)
            removed += 1
        streak.clear()

    for child in list(body.iterchildren()):
        tag = child.tag.split("}")[-1]
        if tag != "p":
            flush()
            continue
        txt = paragraph_text(child)
        if txt == "" and not has_image(child) and not has_page_break(child):
            if get_style(child) in (HEADING1, HEADING2, HEADING3):
                flush()
                continue
            streak.append(child)
        else:
            flush()
    flush()
    return removed


def justify_body(body):
    count = 0
    for child in list(body.iterchildren()):
        if child.tag.split("}")[-1] != "p":
            continue
        style = get_style(child)
        if style in (HEADING1, HEADING2, HEADING3, CAPTION_STYLE):
            continue
        if not paragraph_text(child):
            continue
        set_jc(child, "both")
        count += 1
    return count


def set_hanging_indent(p_elem, twips=720):
    pPr = p_elem.find(w_tag("pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    ind = pPr.find(w_tag("ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(w_tag("left"), str(twips))
    ind.set(w_tag("hanging"), str(twips))


def _run_text(r):
    return "".join(x.text or "" for x in r.findall(w_tag("t")))


def _unbold_run(r):
    rPr = r.find(w_tag("rPr"))
    if rPr is None:
        return False
    removed = False
    for tag in ("b", "bCs"):
        el = rPr.find(w_tag(tag))
        if el is not None:
            rPr.remove(el)
            removed = True
    return removed


def _split_run_at(r, k):
    """Coupe le run r à l'offset k (dans son texte). Retourne la partie droite."""
    ts = r.findall(w_tag("t"))
    full = "".join(x.text or "" for x in ts)
    left, right = full[:k], full[k:]
    for x in ts:
        r.remove(x)
    if left:
        tl = OxmlElement("w:t")
        tl.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        tl.text = left
        r.append(tl)
    r2 = OxmlElement("w:r")
    rPr = r.find(w_tag("rPr"))
    if rPr is not None:
        r2.append(copy.deepcopy(rPr))
    if right:
        tr = OxmlElement("w:t")
        tr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        tr.text = right
        r2.append(tr)
    r.addnext(r2)
    return r2


def _apply_term(p_elem, rx, num, want_label):
    runs = p_elem.findall(w_tag("r"))
    segments = []
    pos = 0
    for r in runs:
        t = _run_text(r)
        segments.append((r, pos, pos + len(t), len(t)))
        pos += len(t)
    full = "".join(_run_text(r) for r in runs)
    matches = list(rx.finditer(full))
    if not matches:
        return 0, False

    unbolded = 0
    for m in matches:
        s, e = m.start(), m.end()
        for (r, rs, re_, _ln) in segments:
            if rs < e and re_ > s:
                if _unbold_run(r):
                    unbolded += 1

    did_label = False
    if want_label:
        e = matches[0].end()
        for (r, rs, re_, ln) in segments:
            if rs < e <= re_:
                k = e - rs
                if k < ln:
                    _split_run_at(r, k)
                r.addnext(make_superscript_run(num))
                did_label = True
                break
    return unbolded, did_label


def place_keyword_labels(body):
    placed = set()
    patterns = []
    for num, _term, pat in GLOSSARY_TERMS:
        rx = re.compile(
            r"(?<![0-9A-Za-zÀ-ÿ])(?:" + pat + r")(?![0-9A-Za-zÀ-ÿ])", re.IGNORECASE
        )
        patterns.append((num, rx))

    labels = 0
    unbold = 0
    for child in list(body.iterchildren()):
        if child.tag.split("}")[-1] != "p":
            continue
        style = get_style(child)
        up = paragraph_text(child).upper().replace("\xa0", " ")
        if style == HEADING1 and up.startswith("GLOSSAIRE"):
            break
        if style in (HEADING1, HEADING2, HEADING3, CAPTION_STYLE):
            continue
        for num, rx in patterns:
            u, did = _apply_term(child, rx, num, num not in placed)
            unbold += u
            if did:
                placed.add(num)
                labels += 1
    return {"labels": labels, "unbold": unbold, "placed": sorted(placed)}


def number_glossary(doc):
    if not doc.tables:
        return 0
    t = doc.tables[-1]
    n = 0
    for i, row in enumerate(t.rows[1:], start=1):
        para = row.cells[0].paragraphs[0]
        if para.runs:
            para.runs[0].text = f"{i}. " + para.runs[0].text
        else:
            para.add_run(f"{i}. ")
        n += 1
    return n


def format_bibliography(body):
    in_region = False
    count = 0
    for child in list(body.iterchildren()):
        if child.tag.split("}")[-1] != "p":
            continue
        style = get_style(child)
        txt = paragraph_text(child)
        up = txt.upper().replace("\xa0", " ")
        if style == HEADING1 and up.startswith("BIBLIOGRAPHIE"):
            in_region = True
            continue
        if style == HEADING1 and up.startswith("GLOSSAIRE"):
            break
        if in_region and txt:
            if re.search(r"\((?:19|20)\d{2}", txt) or "http" in txt.lower():
                set_jc(child, "both")
                set_hanging_indent(child)
                count += 1
    return count


def add_page_numbers(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        para = footer.paragraphs[0]
        for r in list(para._element.findall(w_tag("r"))):
            para._element.remove(r)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        el = para._element
    else:
        el = OxmlElement("w:p")
        footer._element.append(el)
        set_jc(el, "center")

    el.append(make_run("Page "))
    for r in make_field("PAGE", "1"):
        el.append(r)
    el.append(make_run(" / "))
    for r in make_field("NUMPAGES", "1"):
        el.append(r)


def set_document_update_fields(doc):
    settings = doc.settings.element
    update = settings.find(w_tag("updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(w_tag("val"), "true")


def format_margins(section):
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def _renumber_heading(par, old, new):
    runs = par.runs
    full = "".join(r.text for r in runs)
    if old not in full:
        return False
    full = full.replace(old, new, 1)
    if runs:
        runs[0].text = full
        for r in runs[1:]:
            r.text = ""
    else:
        par.add_run(full)
    return True


def insert_app_screenshots(doc, body):
    """Insère une section « Présentation des interfaces réalisées » avec les
    captures d'écran, juste avant la conclusion du Chapitre 4."""
    concl = None
    for p in doc.paragraphs:
        if get_style(p._element) == HEADING2 and p.text.strip().lower().startswith(
            "4.11 conclusion"
        ):
            concl = p
            break
    if concl is None:
        print("   ! Conclusion du Chapitre 4 introuvable, captures non inserees.")
        return 0

    # La conclusion devient 4.12 ; la nouvelle section prend le numero 4.11
    _renumber_heading(concl, "4.11", "4.12")
    concl_p = concl._element

    def add_text(text, style=NORMAL, jc="both", bold=False, italic=False):
        el = OxmlElement("w:p")
        set_style(el, style)
        if jc:
            set_jc(el, jc)
        el.append(make_run(text, bold=bold, italic=italic))
        concl_p.addprevious(el)

    add_text("4.11 Présentation des interfaces réalisées", style=HEADING2, jc=None)
    add_text(
        "Cette section présente les principales interfaces de la plateforme "
        "BlockTask telle qu'elle a été réalisée. Les captures sont organisées "
        "par espace utilisateur — visiteur, client, prestataire et "
        "administrateur — afin d'illustrer le parcours complet décrit dans les "
        "sections précédentes.",
        style=NORMAL,
        jc="both",
    )

    count = 0
    placeholder = 5  # le Chapitre 4 compte deja 4 figures (4.1 a 4.4)
    for item in SCREENSHOTS:
        if item[0] == "group":
            add_text(item[1], style=NORMAL, jc=None, bold=True)
            continue
        _, suffix, caption = item
        path = ASSET_DIR / (SS_PREFIX + suffix)
        if not path.exists():
            print(f"   ! Image manquante : {path.name}")
            continue
        img_par = concl.insert_paragraph_before()
        img_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_par.add_run().add_picture(str(path), width=Cm(15))
        cap = make_caption_paragraph("Figure", "Figure", caption, "4", placeholder)
        concl_p.addprevious(cap)
        placeholder += 1
        count += 1
    return count


def main():
    print(f"Copie : {DST.name}")
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))
    body = doc.element.body

    print("1/9 - Styles de titres + sauts de page...")
    h = fix_heading_styles(body)
    print("2/9 - Legendes numerotees par chapitre + paragraphes explicatifs...")
    c = process_captions(body)
    print("3/9 - Mots-cles : etiquettes chiffrees en exposant...")
    kw = place_keyword_labels(body)
    print("4/9 - Bibliographie APA (retrait suspendu)...")
    bib = format_bibliography(body)
    print("5/9 - Listes automatiques figures/tableaux...")
    lists = replace_manual_lists(body)
    print("6/9 - Suppression des pages blanches...")
    removed = collapse_blanks(body, keep=1)
    print("7/9 - Justification du texte...")
    just = justify_body(body)
    print("7b/9 - Insertion des captures d'application (Chapitre 4)...")
    ss = insert_app_screenshots(doc, body)
    print("8/9 - Numerotation du glossaire + pagination + marges...")
    gl = number_glossary(doc)
    for section in doc.sections:
        format_margins(section)
        add_page_numbers(section)
    print("9/9 - Actualisation des champs a l'ouverture...")
    set_document_update_fields(doc)

    doc.save(str(DST))

    print("\n=== TERMINE ===")
    print(f"Fichier : {DST}")
    print(f"Titres : Titre1={h['h1']}  Titre2={h['h2']}  Titre3={h['h3']}")
    print(f"Sauts de page chapitres : {h['pb']}")
    print(f"Legendes : {c['figures']} figures, {c['tables']} tableaux (+ paragraphes explicatifs)")
    print(f"Mots-cles etiquetes : {kw['labels']} (numeros places: {kw['placed']}), occurrences degraissees: {kw['unbold']}")
    print(f"References APA formatees : {bib}")
    print(f"Entrees glossaire numerotees : {gl}")
    print(f"Captures d'application inserees : {ss} figures (Chapitre 4)")
    print(f"Listes manuelles remplacees : {lists}")
    print(f"Paragraphes vides supprimes : {removed}")
    print(f"Paragraphes justifies : {just}")
    print("\nDans Word : Ctrl+A puis F9 pour actualiser numeros, listes et pagination.")


if __name__ == "__main__":
    main()
